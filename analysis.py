import os
import sys
import json
from datetime import datetime
from pathlib import Path

# 修正 Windows 終端編碼問題
sys.stdout.reconfigure(encoding='utf-8')

# ============= 路徑設定 =============
WORKSPACE = r"C:\Users\garde\OneDrive\3.教務\1.中共軍事體制研究\每日資料蒐集更新"
ENV_FILE  = Path(WORKSPACE) / ".env"

# ── 從 .env 讀取設定 ──────────────────────────────────────
def _load_env() -> dict:
    config = {}
    if ENV_FILE.exists():
        for line in ENV_FILE.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                config[k.strip()] = v.strip()
    return config

_cfg = _load_env()

ANTHROPIC_KEY = _cfg.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_API_KEY", "")
GROK_KEY      = _cfg.get("GROK_API_KEY")      or os.environ.get("GROK_API_KEY", "")

# ── 自動選擇 API 後端 ─────────────────────────────────────
#
#   優先順序：Claude (Anthropic) > Grok (x.ai)
#   - ANTHROPIC_API_KEY 有值 → 使用 Claude
#   - 否則，GROK_API_KEY 有值 → 使用 Grok
#   - 兩者皆無 → 結束並提示
#
if ANTHROPIC_KEY:
    API_BACKEND = "anthropic"
    try:
        import anthropic as _anthropic_sdk
        client = _anthropic_sdk.Anthropic(api_key=ANTHROPIC_KEY)
    except ImportError:
        print("❌ 請安裝 anthropic 套件：pip install anthropic")
        sys.exit(1)
    print(f"✅ 使用 Claude (Anthropic)  — Key 前10碼：{ANTHROPIC_KEY[:10]}...")

elif GROK_KEY:
    API_BACKEND = "grok"
    try:
        from openai import OpenAI as _OpenAI
        client = _OpenAI(api_key=GROK_KEY, base_url="https://api.x.ai/v1")
    except ImportError:
        print("❌ 請安裝 openai 套件：pip install openai")
        sys.exit(1)
    print(f"✅ 使用 Grok (x.ai)         — Key 前10碼：{GROK_KEY[:10]}...")

else:
    print("❌ 找不到任何可用的 API Key！")
    print("   請在 .env 設定 ANTHROPIC_API_KEY 或 GROK_API_KEY")
    sys.exit(1)

# 建立資料夾路徑
INPUT_DIR   = Path(WORKSPACE) / "_daily_input"
OUTPUT_DIR  = Path(WORKSPACE) / "_daily_output"
ARCHIVE_DIR = Path(WORKSPACE) / "_archive"

# ============= 初始化 =============
os.makedirs(INPUT_DIR,   exist_ok=True)
os.makedirs(OUTPUT_DIR,  exist_ok=True)
os.makedirs(ARCHIVE_DIR, exist_ok=True)


# ============= Wiki 知識庫上下文 =============
WIKI_DIR = Path(WORKSPACE) / "wiki"

def load_wiki_context(x_content: str) -> str:
    """從 wiki 知識庫載入相關上下文，注入分析 prompt"""
    context_parts = []
    
    # 1. 讀取索引（快速概覽已知知識）
    index_file = WIKI_DIR / "_index.md"
    if index_file.exists():
        idx_text = index_file.read_text(encoding="utf-8")
        # 只取概念索引表格部分（節省 token）
        if "## 📌 概念文章索引" in idx_text:
            start = idx_text.index("## 📌 概念文章索引")
            end = idx_text.index("---", start + 10) if "---" in idx_text[start + 10:] else start + 2000
            context_parts.append(idx_text[start:start + end - start][:1500])
    
    # 2. 掃描 concepts/ 找出與今日情資相關的概念文章
    concepts_dir = WIKI_DIR / "concepts"
    if concepts_dir.exists():
        for concept_file in concepts_dir.glob("*.md"):
            concept_name = concept_file.stem
            # 簡易相關性判斷：概念名稱的關鍵字是否出現在今日情資中
            keywords = concept_name.replace("與", " ").replace("及", " ").replace("和", " ").split()
            if any(kw in x_content for kw in keywords if len(kw) >= 2):
                try:
                    content = concept_file.read_text(encoding="utf-8")
                    # 取前 800 字元作為上下文
                    context_parts.append(f"--- 已知概念：{concept_name} ---\n{content[:400]}")
                except:
                    pass

    if context_parts:
        return "\n\n".join(context_parts[:5])  # 最多 5 個上下文區塊（節省 token）
    return ""


# ============= 核心分析函數 =============
def analyze_x_post(x_content: str) -> dict:
    """
    分析情資文件，產出：
    1. 原始摘要（中英對照）
    2. 核心分析 + 歷史案例比較
    3. 政策建議
    4. Chicago Style 引用與參考文獻

    自動依 API_BACKEND 切換 Claude / Grok。
    分析前會自動注入 wiki 知識庫上下文。
    """

    # 注入 wiki 上下文
    wiki_context = load_wiki_context(x_content)
    wiki_section = ""
    if wiki_context:
        wiki_section = f"""
⚠️ 以下是知識庫中與今日情資相關的已知背景知識。請在分析時參考這些過往資訊，
特別是在「核心分析」的「歷史案例與關聯性」段落中，引用這些已知的持續追蹤議題：

{wiki_context}

========================================
"""
        print(f"  📚 已注入 wiki 上下文（{len(wiki_context)} 字元）")

    prompt = f"""你是台灣國防與安全研究的資深專家分析師，專長為中共軍事戰略、衛星影像標註分析、灰色地帶作戰與印太安全。
{wiki_section}
今日從各管道蒐集到以下中共軍事動態相關情資（由特殊分隔線 --- 分開）：

========================================
{x_content}
========================================

### 輸出格式範本限制如下（請完全依循此層級與標題）：

# 綜合分析報告

## 一、執行摘要（Executive Summary）
（在此提供 150 字的綜合快速結語）

## 二、當日重點事件摘要
（將上述情資拆解，若沒有網址或來源則寫出「未提供」。必須使用清單列出）

### 事件一：[事件名稱]
- **來源**：[來源/媒體機構]
- **英文原文**：[若有英文原文關鍵句則附上，若無則填寫未提供]
- **當日數據**：[若有具體數量或關鍵指標則列出]
- **中文摘要**：[事件的精簡中文描述]
- **參考連結**：[請確保是標準 Markdown 連結格式。⚠️ 絕對不可捏造、發明或臆測任何 URL（如 bit.ly）]

### 事件二：[事件名稱]
...

## 三、核心分析
（將分析再分為幾個子節，探討戰略與影響）

### 3.1 戰略意圖與對台潛在影響
（論述重點）
### 3.2 中共軍事發展歷史案例與關聯性
（必須包含！請對比過往歷史案例與當前軍事發展的關聯）

## 四、政策建議
（提供針對國防或民防的建議）

## 五、引用文獻（Chicago Style）
（建立正式引用列表，遵守 Chicago Style）

## 六、相關議題深度分析清單
（列出未來值得深究的子議題）
"""

    system_msg = "你是台灣國防與安全研究的資深專家分析師，專長為中共軍事戰略、衛星影像標註分析、灰色地帶作戰與印太安全。"

    # ── Claude (Anthropic) ────────────────────────────────
    if API_BACKEND == "anthropic":
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            system=system_msg,
            messages=[{"role": "user", "content": prompt}]
        )
        text = message.content[0].text

    # ── Grok (x.ai) ───────────────────────────────────────
    else:
        response = client.chat.completions.create(
            model="grok-3",
            max_tokens=3000,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": prompt}
            ]
        )
        text = response.choices[0].message.content

    return {
        "analysis":  text,
        "timestamp": datetime.now().isoformat(),
        "backend":   API_BACKEND,
    }


# ============= 主程序 =============
def main():
    # Step 0: 嘗試自動讀取 Gmail（Grok 郵件）
    try:
        from gmail_reader import GrokEmailReader
        reader = GrokEmailReader()
        gmail_files = reader.run()
        if gmail_files:
            print(f"📬 Gmail 自動讀取: {len(gmail_files)} 封新郵件")
    except FileNotFoundError:
        print("⚠️ Gmail 認證未設定，使用手動輸入模式")
    except ImportError:
        print("⚠️ gmail_reader 模組未安裝，使用手動輸入模式")
    except Exception as e:
        print(f"⚠️ Gmail 讀取失敗，使用手動輸入模式: {e}")

    # 掃描輸入資料夾（含手動 + Gmail 自動讀取的檔案，支援 txt/pdf/docx）
    input_files = []
    input_files.extend(list(INPUT_DIR.glob("*.txt")))
    input_files.extend(list(INPUT_DIR.glob("*.pdf")))
    input_files.extend(list(INPUT_DIR.glob("*.docx")))

    if not input_files:
        print("❌ 沒有找到輸入檔案")
        return

    print(f"✅ 找到 {len(input_files)} 個輸入檔案，開始合併讀取")
    all_x_content = ""
    processed_count = 0

    for input_file in input_files:
        print(f"\n📄 嘗試讀取: {input_file.name}")
        file_text = ""
        try:
            if input_file.suffix.lower() == '.pdf':
                try:
                    import pypdf
                    reader = pypdf.PdfReader(str(input_file))
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            file_text += extracted + "\n"
                except Exception as e:
                    print(f"⚠️ 解析 PDF 失敗: {e}")
            elif input_file.suffix.lower() in ['.docx', '.doc']:
                try:
                    import document
                    # 我們將稍後獨立處理 Doc/Docx (Spire.Doc) 這邊先使用簡單備用語法
                    from spire.doc import Document
                    doc = Document()
                    doc.LoadFromFile(str(input_file))
                    file_text = doc.GetText()
                    # 移除 Spire.Doc 浮水印
                    file_text = file_text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "").strip()
                except ImportError:
                    try:
                        import docx
                        doc = docx.Document(str(input_file))
                        file_text = "\n".join([p.text for p in doc.paragraphs])
                    except Exception as e2:
                        print(f"⚠️ 解析 DOCX/DOC 失敗: {e2}")
                except Exception as e:
                    print(f"⚠️ 解析 DOCX/DOC 失敗: {e}")
            else:
                with open(input_file, 'r', encoding='utf-8') as f:
                    file_text = f.read()
        except Exception as e:
            print(f"⚠️ 讀取檔案 {input_file.name} 內部錯誤: {e}")
            continue

        if not file_text.strip():
            print(f"⚠️ {input_file.name} 擷取為空，跳過。")
            continue

        all_x_content += f"\n\n--- 來源：{input_file.name} ---\n{file_text}\n"
        processed_count += 1
        # 先不要馬上移動，等最後分析成功再統一歸檔
    
    if processed_count == 0 or not all_x_content.strip():
        print("❌ 沒有成功擷取到任何文字，結束程序。")
        return

    print("\n大整合完畢，開始進行綜合分析...")
    generated_report = None
    
    try:
        result = analyze_x_post(all_x_content)
        analysis = result["analysis"]

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = OUTPUT_DIR / f"{timestamp}_綜合分析報告.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(analysis)

        print(f"✅ 已生成綜合分析報告: {output_file.name}")
        generated_report = output_file

        # 統一歸檔所有輸入檔案
        for input_file in input_files:
            archive_file = ARCHIVE_DIR / f"{timestamp}_{input_file.name}"
            try:
                input_file.rename(archive_file)
                print(f"📁 已歸檔: {archive_file.name}")
            except Exception as e:
                pass

    except Exception as e:
        print(f"❌ 總和分析失敗: {e}")

    # 發送通知
    if generated_report:
        try:
            from notifier import Notifier
            notifier = Notifier()
            print(f"\n📤 發送總結通知...")
            res = notifier.notify_report(generated_report)
            if res['email']:
                print(f"  📧 Email 已發送")
            if res['line']:
                print(f"  📱 LINE 已推送")
            print("✅ 通知發送完成")
        except Exception as e:
            print(f"⚠️ 通知發送失敗: {e}")
            
        print("\n🧠 開始編譯情報知識庫 (Wiki) 與派發全團隊審查 (Pull Request)...")
        import subprocess
        try:
            subprocess.run([sys.executable, "wiki_compiler.py", "--team-sync"], check=True)
        except Exception as e:
            print(f"⚠️ 知識庫更新與 PR 發送失敗: {e}")

if __name__ == "__main__":
    main()
